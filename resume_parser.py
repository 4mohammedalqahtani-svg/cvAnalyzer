import io

from pypdf import PdfReader
from docx import Document


def _parse_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages_text = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages_text.append(text)
    return "\n".join(pages_text).strip()


def _parse_docx(file_bytes: bytes) -> str:
    document = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs).strip()


def _parse_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "latin-1"):
        try:
            return file_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore").strip()


def extract_resume_text(filename: str, file_bytes: bytes) -> str:
    lower_name = filename.lower()
    if lower_name.endswith(".pdf"):
        return _parse_pdf(file_bytes)
    if lower_name.endswith(".docx"):
        return _parse_docx(file_bytes)
    if lower_name.endswith(".txt"):
        return _parse_txt(file_bytes)
    raise ValueError(f"Unsupported file type: {filename}. Please upload a PDF, DOCX, or TXT file.")
